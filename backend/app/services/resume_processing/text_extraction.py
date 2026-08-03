import base64
import io
import logging
import pdfplumber
import pypdf
import docx

logger = logging.getLogger(__name__)


def extract_text_pdf(file_bytes: bytes) -> tuple[str, str]:
    """
    Extracts plain text from PDF bytes using a multi-strategy fallback pipeline.

    Returns:
        (text, method) where method is one of:
          - "text_layer"  — fast path: embedded text layer found by pdfplumber/pypdf/pymupdf
          - "ocr"         — slow path: Tesseract OCR (for scanned / image-only PDFs)
          - "vision_ocr"  — fallback: Gemini Vision API (when Tesseract is unavailable)

    Callers should persist `method` as `extraction_method` on `resume_files` so that
    degraded-but-successful OCR paths are visible rather than silent (mirrors the
    `parser_version` pattern already used for LLM vs. heuristic parsing).

    Pipeline order:
    1. pdfplumber standard/layout/word extraction.
    2. pypdf.PdfReader stream extraction.
    3. PyMuPDF embedded text extraction.
    4. Tesseract OCR via pdf2image + pytesseract (scanned PDFs, requires system binaries).
    5. Gemini Vision OCR (final fallback when Tesseract is not installed).
    """
    # -------------------------------------------------------------------------
    # Tier 1-3: text-layer extraction (fast path — try these first on every PDF)
    # -------------------------------------------------------------------------
    text_chunks: list[str] = []

    # Tier 1-3: pdfplumber extraction under a single open handle
    try:
        with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
            # Tier 1: Standard
            for page in pdf.pages:
                pt = page.extract_text()
                if pt and pt.strip():
                    text_chunks.append(pt.strip())

            # Tier 2: Layout mode if standard extracted nothing
            if not text_chunks:
                for page in pdf.pages:
                    pt = page.extract_text(layout=True)
                    if pt and pt.strip():
                        text_chunks.append(pt.strip())

            # Tier 3: Word token extraction if layout mode extracted nothing
            if not text_chunks:
                for page in pdf.pages:
                    words = page.extract_words()
                    if words:
                        page_str = " ".join(w.get("text", "") for w in words if w.get("text"))
                        if page_str.strip():
                            text_chunks.append(page_str.strip())
    except Exception as e:
        logger.warning(f"pdfplumber extraction failed: {e}")

    result = "\n".join(text_chunks).strip()
    if result:
        return result, "text_layer"

    # Tier 4: pypdf.PdfReader fallback
    text_chunks = []
    try:
        reader = pypdf.PdfReader(io.BytesIO(file_bytes))
        for page in reader.pages:
            pt = page.extract_text()
            if pt and pt.strip():
                text_chunks.append(pt.strip())
    except Exception as e:
        logger.warning(f"pypdf extraction failed: {e}")

    result = "\n".join(text_chunks).strip()
    if result:
        return result, "text_layer"

    # Tier 5: PyMuPDF embedded text extraction
    text_chunks = []
    try:
        import fitz  # pymupdf
        doc = fitz.open(stream=file_bytes, filetype="pdf")
        for page in doc:
            pt = page.get_text("text")
            if pt and pt.strip():
                text_chunks.append(pt.strip())
        doc.close()
    except Exception as e:
        logger.warning(f"PyMuPDF text extraction failed: {e}")

    result = "\n".join(text_chunks).strip()
    if result:
        return result, "text_layer"

    # -------------------------------------------------------------------------
    # Tier 6: Tesseract OCR — for scanned / image-only PDFs
    # Requires: tesseract-ocr and poppler-utils system packages (installed in
    # Dockerfile). Only reached when all text-layer tiers return empty.
    # -------------------------------------------------------------------------
    logger.info("No text layer found in PDF; falling back to OCR...")
    ocr_text = _extract_text_ocr(file_bytes)
    if ocr_text:
        return ocr_text, "ocr"

    # -------------------------------------------------------------------------
    # Tier 7: Gemini Vision OCR — final fallback when Tesseract is not available
    # (e.g. local dev without system binaries). Renders PDF pages as PNG images
    # and sends them to Gemini Vision API. No binary dependencies required.
    # -------------------------------------------------------------------------
    logger.info("Tesseract OCR returned empty or unavailable; trying Gemini Vision OCR...")
    vision_text = _extract_text_via_gemini_vision(file_bytes)
    if vision_text:
        return vision_text, "vision_ocr"

    return "", "text_layer"


def _extract_text_ocr(file_bytes: bytes) -> str:
    """
    Rasterises each PDF page with pdf2image (poppler backend) then runs
    pytesseract over each image. Only invoked when all text-layer tiers fail.

    Requires system packages: `tesseract-ocr` (OCR engine) and `poppler-utils`
    (pdftoppm, used by pdf2image). Both are added to the Dockerfile.
    """
    from app.config import settings

    try:
        import pytesseract
        from pdf2image import convert_from_bytes
    except ImportError:
        logger.warning("pytesseract/pdf2image not installed. Skipping Tesseract OCR.")
        return ""

    try:
        images = convert_from_bytes(file_bytes, dpi=200)
    except Exception as e:
        logger.warning(f"pdf2image failed to rasterise PDF (poppler may be missing): {e}")
        return ""

    if len(images) > settings.ocr_max_pages:
        logger.warning(
            f"PDF has {len(images)} pages, exceeding OCR_MAX_PAGES={settings.ocr_max_pages}. "
            f"Only OCR-ing the first {settings.ocr_max_pages} pages."
        )
        images = images[: settings.ocr_max_pages]

    text_chunks: list[str] = []
    for i, image in enumerate(images):
        try:
            page_text = pytesseract.image_to_string(image, timeout=settings.ocr_timeout_seconds)
            if page_text.strip():
                text_chunks.append(page_text.strip())
        except RuntimeError as e:
            # pytesseract raises RuntimeError specifically on timeout
            logger.warning(f"Tesseract OCR timed out on page {i + 1}: {e}")
        except Exception as e:
            logger.warning(f"Tesseract OCR failed on page {i + 1}: {e}")

    result = "\n".join(text_chunks).strip()
    if result:
        logger.info(f"Tesseract OCR succeeded: {len(text_chunks)} page(s) extracted.")
    return result


def _extract_text_via_gemini_vision(file_bytes: bytes) -> str:
    """
    Renders each PDF page to a PNG image using PyMuPDF and sends the images
    to Gemini Vision API (gemini-2.0-flash) for full text transcription.

    This is the last-resort fallback for environments where Tesseract is not
    installed (e.g. local dev on Windows). In Docker production the Tesseract
    tier should succeed before this is ever reached.

    Returns the concatenated raw text across all pages, or empty string on failure.
    """
    try:
        import fitz  # pymupdf
        import httpx
        from app.config import settings

        api_key = getattr(settings, "gemini_api_key", None)
        if not api_key:
            logger.warning("Gemini Vision OCR: GEMINI_API_KEY not configured. Skipping.")
            return ""

        ocr_max_pages = getattr(settings, "ocr_max_pages", 5)
        doc = fitz.open(stream=file_bytes, filetype="pdf")
        page_count = len(doc)
        logger.info(f"Gemini Vision OCR: rendering {min(page_count, ocr_max_pages)} page(s) at 150 DPI...")

        all_text: list[str] = []
        model = getattr(settings, "gemini_model", "") or "gemini-2.0-flash"
        url = f"https://generativelanguage.googleapis.com/v1beta/models/{model}:generateContent?key={api_key}"

        for page_num in range(min(page_count, ocr_max_pages)):
            page = doc[page_num]
            pix = page.get_pixmap(dpi=150)
            png_bytes = pix.tobytes("png")
            b64_image = base64.b64encode(png_bytes).decode("utf-8")

            payload = {
                "contents": [
                    {
                        "parts": [
                            {
                                "text": (
                                    "You are a professional document transcriber. "
                                    "Transcribe ALL text visible in this resume page image exactly as written, "
                                    "preserving headings, bullet points, and all content. "
                                    "Output only the raw text — no JSON, no commentary, no markdown wrappers."
                                )
                            },
                            {
                                "inline_data": {
                                    "mime_type": "image/png",
                                    "data": b64_image
                                }
                            }
                        ]
                    }
                ],
                "generationConfig": {
                    "temperature": 0.1,
                    "maxOutputTokens": 4096
                }
            }

            try:
                with httpx.Client(timeout=30.0) as client:
                    resp = client.post(url, json=payload)
                    if resp.status_code == 200:
                        data = resp.json()
                        candidates = data.get("candidates", [])
                        if candidates:
                            parts = candidates[0].get("content", {}).get("parts", [])
                            page_text = " ".join(p.get("text", "") for p in parts if p.get("text"))
                            if page_text.strip():
                                all_text.append(page_text.strip())
                                logger.info(f"Gemini OCR page {page_num + 1}: extracted {len(page_text)} chars")
                    else:
                        logger.warning(
                            f"Gemini Vision OCR page {page_num + 1} HTTP {resp.status_code}: {resp.text[:200]}"
                        )
            except Exception as page_err:
                logger.warning(f"Gemini Vision OCR failed on page {page_num + 1}: {page_err}")

        doc.close()
        result = "\n\n".join(all_text).strip()
        if result:
            logger.info(f"Gemini Vision OCR succeeded. Total characters extracted: {len(result)}")
        else:
            logger.warning("Gemini Vision OCR returned no text.")
        return result

    except Exception as e:
        logger.error(f"Gemini Vision OCR fatal error: {e}", exc_info=True)
        return ""


def extract_text_docx(file_bytes: bytes) -> str:
    """
    Extracts plain text from DOCX bytes using python-docx.
    Traverses both standard paragraphs and table cells to capture all details.
    DOCX always has a real text layer — no OCR path needed.
    """
    doc = docx.Document(io.BytesIO(file_bytes))
    text = []

    # 1. Extract standard paragraphs
    for para in doc.paragraphs:
        if para.text.strip():
            text.append(para.text)

    # 2. Extract text inside tables (resumes often use tables for layout)
    for table in doc.tables:
        for row in table.rows:
            row_text = []
            for cell in row.cells:
                cell_text = cell.text.strip()
                if cell_text and cell_text not in row_text:
                    row_text.append(cell_text)
            if row_text:
                text.append(" | ".join(row_text))

    return "\n".join(text).strip()
